#!/usr/bin/python3
#
# Copyright (c) 2019, Arista Networks EOS+
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice, this
#   list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the Arista nor the names of its
#   contributors may be used to endorse or promote products derived from
#   this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#

"""Utilities for using PyTest in network testing"""

import sys
import logging
import datetime
import json
import re
import yaml
import docx


logging.basicConfig(level=logging.INFO, filename='vane.log', filemode='w',
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')


class ReportClient:
    """ Creates an instance of the Report Client.
    """

    def __init__(self, test_definition):
        """ Initializes the Test Client

            Args:
                test_definition (str): YAML representation of NRFU tests
        """

        logging.info('Convert yaml data-model to a python data structure')
        self.data_model = self._import_yaml(test_definition)
        logging.info('Internal test data-model initialized with value: '
                     f'{self.data_model}')

        _results_file = self.data_model['parameters']['results_file']
        self._results_data = self._import_yaml(_results_file)
        logging.info(f'Results file data is {self._results_data}')

        self._document = docx.Document()
        self._summary_results = self._compile_test_results()
        logging.info(f"Test Results: {self._summary_results}")
        self._major_section = 1

    def _import_yaml(self, yaml_file):
        """ Import YAML file as python data structure

            Args:
                yaml_file (str): Name of YAML file
        """

        logging.info(f'Opening {yaml_file} for read')
        try:
            with open(yaml_file, 'r') as input_yaml:
                try:
                    yaml_data = yaml.safe_load(input_yaml)
                    logging.info(f'Inputed the following yaml: '
                                 f'{yaml_data}')
                    return yaml_data
                except yaml.YAMLError as err_data:
                    logging.error(f'Error in YAML file. {err_data}')
                    sys.exit(1)
        except OSError as err_data:
            logging.error(f'Defintions file: {yaml_file} not '
                          f'found. {err_data}')
            sys.exit(1)

    def write_result_doc(self):
        """ Create MSFT docx with results
        """

        logging.info('Create MSFT docx with results')
        self._write_title_page()
        self._write_summary_report()
        self._write_detail_report()
        self._document.save('../reports/report-class.docx')

    def _return_date(self):
        """ Genreate a formatted date and return to calling
            function.
        """

        date_obj = datetime.datetime.now()
        format_date = date_obj.strftime("%B %d, %Y %I:%M:%S%p")

        logging.info(f'Returning formatted date: {format_date}')
        return format_date

    def _write_title_page(self):
        """ Write report title page
        """

        logging.info('Create report title page')
        format_date = self._return_date()
        self._document.add_heading('Test Report', 0)
        p = self._document.add_paragraph(f'{format_date}')
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        self._document.add_page_break()

    def _write_summary_report(self):
        """ Write summary reports
        """

        self._document.add_heading(f'{self._major_section}. Test Results Summary', 1)
        self._write_summary_results()
        self._write_dut_summary_results()
        self._write_suite_summary_results()

        self._major_section += 1
        self._document.add_page_break()

    def _write_summary_results(self):
        """ Write summary results section
        """

        logging.info("Create summary results table")
        self._document.add_heading(f'{self._major_section}.1 Summary Results', 2)
        table = self._document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Total Tests'
        hdr_cells[1].text = 'Total Passed'
        hdr_cells[2].text = 'Total Failed'
        hdr_cells[3].text = 'Total Skipped'
        hdr_cells[4].text = 'Total Errored'
        hdr_cells[5].text = 'Total Duration'

        ptr = self._summary_results['summaryResults']
        total_tests = self._totals(ptr, 'num_tests')
        total_pass = self._totals(ptr, 'passed')
        total_fail = self._totals(ptr, 'failed')
        total_skip = self._totals(ptr, 'skipped')
        total_err = self._totals(ptr, 'error')
        total_time = self._totals(ptr, 'duration')

        row_cells = table.add_row().cells
        row_cells[0].text = total_tests
        row_cells[1].text = total_pass
        row_cells[2].text = total_fail
        row_cells[3].text = total_skip
        row_cells[4].text = total_err
        row_cells[5].text = total_time

    def _write_dut_summary_results(self):
        """ Write summary DUT result section
        """

        logging.info("Create DUT summary results table")
        self._document.add_heading(f'{self._major_section }.2 Summary Totals '
                                   'for Devices Under Tests', 2)

        table = self._document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DUT'
        hdr_cells[1].text = 'Total Tests'
        hdr_cells[2].text = 'Total Passed'
        hdr_cells[3].text = 'Total Failed'
        hdr_cells[4].text = 'Total Skipped'
        hdr_cells[5].text = 'Total Errored'

        duts = self._summary_results['duts']

        for dut in duts:
            total_tests = self._totals(dut, 'TOTAL')
            total_pass = self._totals(dut, 'PASS')
            total_fail = self._totals(dut, 'FAIL')
            total_skip = self._totals(dut, 'SKIP')
            total_err = self._totals(dut, 'ERROR')
            dut_name = self._totals(dut, 'name')

            row_cells = table.add_row().cells
            row_cells[0].text = dut_name
            row_cells[1].text = total_tests
            row_cells[2].text = total_pass
            row_cells[3].text = total_fail
            row_cells[4].text = total_skip
            row_cells[5].text = total_err

    def _write_suite_summary_results(self):
        """ Write summary test suite result section
        """

        logging.info("Create Suite summary results table")
        self._document.add_heading(f'{self._major_section }.3 Summary Totals '
                                   'for Test Suites', 2)
        suite_results = self._compile_suite_results()

        table = self._document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Suite'
        hdr_cells[1].text = 'Total Tests'
        hdr_cells[2].text = 'Total Passed'
        hdr_cells[3].text = 'Total Failed'

        for suite_result in suite_results:
            ts_name = suite_result['name']
            logging.info(f'Test suite name is {ts_name}')
            ts_name = ts_name.split('.')[0]
            ts_name = ts_name.split('_')[1].upper()
            logging.info(f'Formatted test suite name is {ts_name}')

            row_cells = table.add_row().cells
            row_cells[0].text = ts_name
            row_cells[1].text = str(suite_result['total_tests'])
            row_cells[2].text = str(suite_result['total_pass'])
            row_cells[3].text = str(suite_result['total_fail'])

    def _compile_test_results(self):
        """ Parse PyTest JSON results and compile:
        """

        json_report = self.data_model['parameters']['json_report']
        json_report = f"{json_report}.json"
        test_results = {}
        logging.info(f'Opening JSON file {json_report} to parse for summary results')

        with open(json_report, 'r') as json_file:
            logging.info(f'Raw json report is {json_file}')
            test_data = json.load(json_file)
            tests = test_data["report"]["tests"]
            logging.info(f'Structured json report is {test_data}')

            summary = test_data["report"]["summary"]
            test_results["summaryResults"] = summary
            logging.info(f'Summary for test cases are {summary}')
            test_results["duts"] = self._parse_testcases(tests)

        return test_results

    def _parse_testcases(self, testcases):
        """ Parse Test cases and return compilation per DUT
        """

        testcases_results = []
        dut_list = []

        for testcase in testcases:
            if re.search('\[.*\]', testcase["name"]):
                dut_name = re.findall('\[.*\]', testcase["name"])[0][1:-1]
                test_result = testcase["outcome"]

                if dut_name not in dut_list:
                    dut_list.append(dut_name)
                    testcases_results.append({})
                    testcases_results[-1]["PASS"] = 0
                    testcases_results[-1]["FAIL"] = 0
                    testcases_results[-1]["SKIP"] = 0
                    testcases_results[-1]["ERROR"] = 0
                    testcases_results[-1]["TOTAL"] = 0

                dut_index = dut_list.index(dut_name)
                testcases_results[dut_index]["name"] = dut_name

                if test_result == "passed":
                    testcases_results[dut_index]["PASS"] += 1
                elif test_result == "failed":
                    testcases_results[dut_index]["FAIL"] += 1
                elif test_result == "skipped":
                    testcases_results[dut_index]["SKIP"] += 1
                elif test_result == "error":
                    testcases_results[dut_index]["ERROR"] += 1

                testcases_results[dut_index]["TOTAL"] += 1

        return testcases_results

    def _totals(self, ptr, ptr_key):
        """ Test for a key in dictionary.  If key exists return key and if key is
            missing return 0

        Args:
            ptr (dict): dictionary to check
            ptr_key (str): key to test if in dict

        Retrun:
            total (str): Value in dictionary
        """

        total = "0"

        if ptr_key in ptr:
            total = str(ptr[ptr_key])

        return total

    def _write_detail_report(self):
        """ Write summary reports
        """

        test_suites = self._results_data['test_suites']

        for test_suite in test_suites:
            self._write_detail_major_section(test_suite)
            minor_section = 1

            for test_case in test_suite['test_cases']:
                self._write_detail_minor_section(test_case, minor_section)
                dut_section = 1

                for dut in test_case['duts']:
                    self._write_dut_minor_section(dut,
                                                  minor_section,
                                                  dut_section)
                    dut_section += 1

                minor_section += 1

            self._major_section += 1

    def _write_detail_major_section(self, test_suite):
        """ Write Detailed majore report section

            Args:
                test_suite (dict): test_suite result data
        """

        logging.info(f'Raw test suite data is {test_suite}')
        ts_name = test_suite['name']
        logging.info(f'Test suite name is {ts_name}')
        ts_name = ts_name.split('.')[0]
        ts_name = ts_name.split('_')[1].upper()
        logging.info(f'Formatted test suite name is {ts_name}')
        self._document.add_heading(f'{self._major_section}. Detailed Test '
                                   f'Suite Results: {ts_name}', 1)

    def _write_detail_minor_section(self, test_case, minor_section):
        """[summary]

        Args:
            test_case (dict): test_case result data
            minor_section (int): minor section number
        """

        logging.info(f'Raw test case data is {test_case}')
        tc_name = test_case['name']
        logging.info(f'Test case name is {tc_name}')
        tc_name = ' '.join(tc_name.split('_'))[:-3].upper()
        logging.info(f'Formattted test case name is {tc_name}')
        self._document.add_heading(f'{self._major_section}.{minor_section} '
                                   f'Test Case: {tc_name}', 2)
        description = test_case['duts'][0]['description']
        p = self._document.add_paragraph(f'Description: {description}')

    def _write_dut_minor_section(self, dut, minor_section, dut_section):
        """[summary]
        Args:
            dut ([type]): [description]
            minor_section ([type]): [description]
            dut_section ([type]): [description]
        """

        logging.info(f'Raw DUT data is {dut}')
        dut_name = dut['dut']
        dut_name = dut_name.upper()
        logging.info(f'DUT name is {dut_name}')
        self._document.add_heading(f'{self._major_section}.{minor_section}.'
                                   f'{dut_section} DUT: {dut_name}', 3)

    def _compile_suite_results(self):
        """ Compile test suite results and return them

            Return:
                suite_results (list): List of compiled test suite data
        """

        test_suites = self._results_data['test_suites']
        suite_results = []

        for test_suite in test_suites:
            suite_result = {}
            suite_result['total_tests'] = 0
            suite_result['total_pass'] = 0
            suite_result['total_fail'] = 0
            suite_result['name'] = test_suite['name']

            suite_name = suite_result['name']
            logging.info('Zeroing test_suite results for test suite: '
                         f'{suite_name} and data: {suite_result}')

            for test_case in test_suite['test_cases']:
                for dut in test_case['duts']:
                    suite_result['total_tests'] += 1

                    if dut['test_result']:
                        suite_result['total_pass'] += 1
                    else:
                        suite_result['total_fail'] += 1

            logging.info(f'Compiled test suite data: {suite_result}')
            suite_results.append(suite_result)

        logging.info(f'Compiled suite results: {suite_results}')
        return suite_results
