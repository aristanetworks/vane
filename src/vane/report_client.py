#!/usr/bin/env python3
#
# Copyright (c) 2019, Arista Networks EOS+
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice, this
#   list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the Arista nor the names of its
#   contributors may be used to endorse or promote products derived from
#   this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#

"""Utilities for using PyTest in network testing"""

import sys
import logging
import datetime
import json
import os
import re
import yaml
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(
    level=logging.INFO,
    filename="vane.log",
    filemode="w",
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)


class ReportClient:
    """Creates an instance of the Report Client."""

    def __init__(self, test_definition):
        """Initializes the Test Client

        Args:
            test_definition (str): YAML representation of NRFU tests
        """

        logging.info("Convert yaml data-model to a python data structure")
        self.data_model = self._import_yaml(test_definition)
        logging.info(
            "Internal test data-model initialized with value: "
            f"{self.data_model}"
        )
        self._summary_results = self._compile_test_results()
        logging.info(f"Test Results: {self._summary_results}")

        self._reports_dir = self.data_model["parameters"]["report_dir"]
        _results_dir = self.data_model["parameters"]["results_dir"]
        _results_file = self.data_model["parameters"]["results_file"]
        self._results_datamodel = None
        self._compile_yaml_data(_results_dir, _results_file)
        logging.info(f"Results file data is {self._results_datamodel}")

        self._document = docx.Document()
        self._major_section = 1
        self._test_id = 1

    def _compile_yaml_data(self, yaml_dir, yaml_file):
        """[summary]

        Args:
            yaml_dir ([type]): [description]
            yaml_file ([type]): [description]
        """

        logging.info(
            f"yaml directory is {yaml_dir}\n"
            f"yaml output file is {yaml_file}")
        yaml_files = os.listdir(yaml_dir)
        logging.info(f"yaml input files are {yaml_files}")

        for name in yaml_files:
            if "result-" in name:
                yaml_file = f"{yaml_dir}/{name}"
                yaml_data = self._import_yaml(yaml_file)

                self._reconcile_results(yaml_data)
            else:
                logging.error(f"Incorrect filename: {name}")

        logging.info(f"Updated results_data to {self._results_datamodel}")

    def _reconcile_results(self, test_parameters):
        """[summary]

        Args:
            yaml_data ([type]): [description]
        """

        test_suite = test_parameters["test_suite"]
        test_suite = test_suite.split("/")[-1]
        dut_names = test_parameters["duts"]
        test_case = test_parameters["name"]
        yaml_data = self._results_datamodel

        if not self._results_datamodel:
            self._results_datamodel = {
                "test_suites": [
                    {
                        "name": test_suite,
                        "test_cases": [{"name": test_case, "duts": []}],
                    }
                ]
            }

        logging.info(
            f"\n\n\rFind the Index for test suite: {test_suite} on dut "
            f"{dut_names}\n\n\r"
        )
        logging.info(self._results_datamodel["test_suites"])
        test_suites = [
            param["name"] for param in self._results_datamodel["test_suites"]
        ]

        if test_suite in test_suites:
            suite_index = test_suites.index(test_suite)
            logging.info(
                f"Test suite {test_suite} exists in results file at "
                f"index {suite_index}"
            )
        else:
            logging.info(f"Create test suite {test_suite} in results file")
            suite_stub = {"name": test_suite, "test_cases": []}
            self._results_datamodel["test_suites"].append(suite_stub)
            suite_index = len(self._results_datamodel["test_suites"]) - 1

        logging.info(
            f"Find Index for test case: {test_case} on duts {dut_names}")
        test_cases = [
            param["name"]
            for param in self._results_datamodel["test_suites"][suite_index][
                "test_cases"
            ]
        ]

        if test_case in test_cases:
            test_index = test_cases.index(test_case)
            logging.info(
                f"Test case {test_case} exists in results file at index "
                f"{test_index}"
            )
        else:
            logging.info(f"Create test case {test_case} in results file")
            test_stub = {"name": test_case, "duts": []}
            self._results_datamodel["test_suites"][suite_index][
                "test_cases"
            ].append(test_stub)
            test_index = (
                len(
                    self._results_datamodel["test_suites"][suite_index][
                        "test_cases"
                    ]
                )
                - 1
            )

        logging.info(f"Find Index for duts {dut_names}")
        duts = [
            param["duts"]
            for param in self._results_datamodel["test_suites"][suite_index][
                "test_cases"
            ][test_index]["duts"]
        ]

        if dut_names not in duts:
            logging.info(
                f"Add DUT/s {dut_names} to test case {test_case} with "
                f"parameters {test_parameters}"
            )
            yaml_ptr = self._results_datamodel["test_suites"][suite_index]
            yaml_ptr["test_cases"][test_index]["duts"].append(test_parameters)

    def _import_yaml(self, yaml_file):
        """Import YAML file as python data structure

        Args:
            yaml_file (str): Name of YAML file
        """

        logging.info(f"Opening {yaml_file} for read")
        try:
            with open(yaml_file, "r") as input_yaml:
                try:
                    yaml_data = yaml.safe_load(input_yaml)
                    logging.info(
                        f"Inputed the following yaml: "
                        f"{yaml_data}")
                    return yaml_data
                except yaml.YAMLError as err_data:
                    logging.error(f"Error in YAML file. {err_data}")
                    sys.exit(1)
        except OSError as err_data:
            logging.error(
                f"Defintions file: {yaml_file} not " f"found. {err_data}"
            )
            sys.exit(1)

    def write_result_doc(self):
        """Create MSFT docx with results"""

        logging.info("Create MSFT docx with results")
        self._write_title_page()
        self._write_toc_page()
        self._write_summary_report()
        self._write_tests_case_report()
        self._write_detail_report()

        _, file_date = self._return_date()
        reports_dir = self._reports_dir
        file_name = f"{reports_dir}/report_{file_date}.docx"
        logging.info("Writing docx report to file: %s" % (file_name))
        self._document.save(file_name)

    def _return_date(self):
        """Genreate a formatted date and return to calling
        function.
        """

        date_obj = datetime.datetime.now()
        format_date = date_obj.strftime("%B %d, %Y %I:%M:%S%p")
        file_date = date_obj.strftime("%y%m%d%H%M")

        logging.info(f"Returning formatted date: {format_date}")
        return format_date, file_date

    def _write_title_page(self):
        """Write report title page"""

        logging.info("Create report title page")
        format_date, _ = self._return_date()
        self._document.add_heading("Test Report", 0)
        paragraph = self._document.add_paragraph(f"{format_date}")
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        self._document.add_page_break()

    def _write_toc_page(self):
        """Write table of contents page"""
        paragraph = self._document.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        fld_char3 = OxmlElement("w:t")
        fld_char3.text = "Right-click to update field."
        fld_char2.append(fld_char3)

        fld_char4 = OxmlElement("w:fldChar")
        fld_char4.set(qn("w:fldCharType"), "end")

        r_element = run._r
        r_element.append(fld_char)
        r_element.append(instr_text)
        r_element.append(fld_char2)
        r_element.append(fld_char4)
        p_element = paragraph._p

        self._document.add_page_break()

    def _write_summary_report(self):
        """Write summary reports"""

        self._document.add_heading(
            f"{self._major_section}. Test Results " "Summary", 1
        )
        self._write_summary_results()
        self._write_dut_summary_results()
        self._write_suite_summary_results()

        self._major_section += 1
        self._document.add_page_break()

    def _write_summary_results(self):
        """Write summary results section"""

        logging.info("Create summary results table")
        self._document.add_heading(
            f"{self._major_section}.1 Summary " "Results", 2
        )
        table = self._document.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Total Tests"
        hdr_cells[1].text = "Total Passed"
        hdr_cells[2].text = "Total Failed"
        hdr_cells[3].text = "Total Skipped"
        hdr_cells[4].text = "Total Errored"
        hdr_cells[5].text = "Total Duration"

        ptr = self._summary_results["summaryResults"]
        total_tests = self._totals(ptr, "num_tests")
        total_pass = self._totals(ptr, "passed")
        total_fail = self._totals(ptr, "failed")
        total_skip = self._totals(ptr, "skipped")
        total_err = self._totals(ptr, "error")
        total_time = self._totals(ptr, "duration")

        row_cells = table.add_row().cells
        row_cells[0].text = total_tests
        row_cells[1].text = total_pass
        row_cells[2].text = total_fail
        row_cells[3].text = total_skip
        row_cells[4].text = total_err
        row_cells[5].text = total_time

    def _write_dut_summary_results(self):
        """Write summary DUT result section"""

        duts = self._summary_results["duts"]
        if not duts:
            logging.info("Skipping DUT summary results table")
            return

        logging.info("Create DUT summary results table")
        self._document.add_heading(
            f"{self._major_section }.2 Summary Totals "
            "for Devices Under Tests",
            2,
        )

        table = self._document.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "DUT/s"
        hdr_cells[1].text = "Total Tests"
        hdr_cells[2].text = "Total Passed"
        hdr_cells[3].text = "Total Failed"
        hdr_cells[4].text = "Total Skipped"
        hdr_cells[5].text = "Total Errored"

        for dut in duts:
            total_tests = self._totals(dut, "TOTAL")
            total_pass = self._totals(dut, "PASS")
            total_fail = self._totals(dut, "FAIL")
            total_skip = self._totals(dut, "SKIP")
            total_err = self._totals(dut, "ERROR")
            dut_name = self._totals(dut, "name")

            row_cells = table.add_row().cells
            row_cells[0].text = dut_name
            row_cells[1].text = total_tests
            row_cells[2].text = total_pass
            row_cells[3].text = total_fail
            row_cells[4].text = total_skip
            row_cells[5].text = total_err

    def _write_suite_summary_results(self):
        """Write summary test suite result section"""

        sub_section = 2
        if self._summary_results["duts"]:
            sub_section = 3

        logging.info("Create Suite summary results table")
        self._document.add_heading(
            f"{self._major_section }.{sub_section} Summary Totals " "for Test Suites", 2
        )
        suite_results = self._compile_suite_results()

        table = self._document.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Test Suite"
        hdr_cells[1].text = "Total Tests"
        hdr_cells[2].text = "Total Passed"
        hdr_cells[3].text = "Total Failed"
        hdr_cells[4].text = "Total Skipped"
        if not suite_results:
            logging.info("Skipping the test suite results")
            return
        for suite_result in suite_results:
            ts_name = self._format_ts_name(suite_result["name"])

            row_cells = table.add_row().cells
            row_cells[0].text = ts_name
            row_cells[1].text = str(suite_result["total_tests"])
            row_cells[2].text = str(suite_result["total_pass"])
            row_cells[3].text = str(suite_result["total_fail"])
            row_cells[4].text = str(suite_result["total_skip"])

    def _compile_test_results(self):
        """Parse PyTest JSON results and compile:"""

        json_report = self.data_model["parameters"]["json_report"]
        json_report = f"{json_report}.json"
        test_results = {}
        logging.info(
            f"Opening JSON file {json_report} to parse for summary " "results"
        )

        with open(json_report, "r") as json_file:
            logging.info(f"Raw json report is {json_file}")
            test_data = json.load(json_file)
            tests = test_data["report"]["tests"]
            logging.info(f"Structured json report is {test_data}")

            summary = test_data["report"]["summary"]
            test_results["summaryResults"] = summary
            logging.info(f"Summary for test cases are {summary}")
            test_results["duts"] = self._parse_testcases(tests)

        return test_results

    def _parse_testcases(self, testcases):
        """Parse Test cases and return compilation per DUT"""

        testcases_results = []
        dut_list = []

        for testcase in testcases:
            if re.search(r"\[.*\]", testcase["name"]):
                dut_name = re.findall(r"\[.*\]", testcase["name"])[0][1:-1]
                test_result = testcase["outcome"]

                if dut_name not in dut_list:
                    dut_list.append(dut_name)
                    testcases_results.append({})
                    testcases_results[-1]["PASS"] = 0
                    testcases_results[-1]["FAIL"] = 0
                    testcases_results[-1]["SKIP"] = 0
                    testcases_results[-1]["ERROR"] = 0
                    testcases_results[-1]["TOTAL"] = 0

                dut_index = dut_list.index(dut_name)
                testcases_results[dut_index]["name"] = dut_name

                if test_result == "passed":
                    testcases_results[dut_index]["PASS"] += 1
                elif test_result == "failed":
                    testcases_results[dut_index]["FAIL"] += 1
                elif test_result == "skipped":
                    testcases_results[dut_index]["SKIP"] += 1
                elif test_result == "error":
                    testcases_results[dut_index]["ERROR"] += 1

                testcases_results[dut_index]["TOTAL"] += 1

        return testcases_results

    def _totals(self, ptr, ptr_key):
        """Test for a key in dictionary.  If key exists return key and if key is
            missing return 0

        Args:
            ptr (dict): dictionary to check
            ptr_key (str): key to test if in dict

        Retrun:
            total (str): Value in dictionary
        """

        total = "0"

        if ptr_key in ptr:
            total = str(ptr[ptr_key])

        return total

    def _write_tests_case_report(self):
        """Write summary test case report"""

        self._document.add_heading(
            f"{self._major_section}. " "Test Case Results Summary", 1
        )
        self._major_section += 1

        testcase_results = self._compile_testcase_results()

        table = self._document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        test_num = 1

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Test Id"
        hdr_cells[1].text = "Test Suite"
        hdr_cells[2].text = "Test Case"
        hdr_cells[3].text = "DUT/s"
        hdr_cells[4].text = "Result"
        hdr_cells[5].text = "Failure Reason"
        if not testcase_results:
            logging.info("Skipping the summary testcase report")
            return
        for testcase_result in testcase_results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(test_num)
            row_cells[1].text = str(testcase_result["test_suite"])
            row_cells[2].text = str(testcase_result["test_case"])
            row_cells[3].text = str(testcase_result["dut"])
            row_cells[4].text = str(testcase_result["results"])
            row_cells[5].text = str(testcase_result["fail_reason"])

            test_num += 1

        self._document.add_page_break()

    def _write_detail_report(self):
        """Write detailed test case report"""

        if not self._results_datamodel:
            logging.info("Skipping the detailed testcase report")
            return
        test_suites = self._results_datamodel["test_suites"]

        for test_suite in test_suites:
            self._write_detail_major_section(test_suite)
            minor_section = 1

            for test_case in test_suite["test_cases"]:
                self._write_detail_minor_section(test_case, minor_section)
                dut_section = 1

                for dut in test_case["duts"]:
                    self._write_detail_dut_section(
                        dut, minor_section, dut_section
                    )
                    dut_section += 1

                minor_section += 1

            self._major_section += 1

    def _write_detail_major_section(self, test_suite):
        """Write Detailed majore report section

        Args:
            test_suite (dict): test_suite result data
        """

        ts_name = self._format_ts_name(test_suite["name"])
        self._document.add_heading(
            f"{self._major_section}. Detailed Test "
            f"Suite Results: {ts_name}",
            1,
        )

    def _write_detail_minor_section(self, test_case, minor_section):
        """[summary]

        Args:
            test_case (dict): test_case result data
            minor_section (int): minor section number
        """

        tc_name = self._format_tc_name(test_case["name"])

        self._document.add_heading(
            f"{self._major_section}.{minor_section} "
            f"Test Case: {tc_name}", 2)

    def _write_detail_dut_section(self, dut, minor_section, dut_section):
        """[summary]
        Args:
            dut ([type]): [description]
            minor_section ([type]): [description]
            dut_section ([type]): [description]
        """

        logging.info(f"Raw DUT data is {dut}")
        dut_names = dut["duts"]
        dut_names = ", ".join(dut_name.upper() for dut_name in dut_names)
        logging.info(f"DUT names: {dut_names}")
        self._document.add_heading(
            f"{self._major_section}.{minor_section}."
            f"{dut_section} DUT/s: {dut_names}",
            3,
        )

        table = self._document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Test Parameter"
        hdr_cells[1].text = "Description"

        self._add_dut_table_row("test_id", dut, table)
        self._add_dut_table_row("name", dut, table)
        self._add_dut_table_row("description", dut, table)
        self._add_dut_table_row("dut", dut, table)
        self._add_dut_table_row("show_cmd", dut, table)
        self._add_dut_table_row("expected_output", dut, table)
        self._add_dut_table_row("actual_output", dut, table)
        self._add_dut_table_row("test_result", dut, table)
        self._add_dut_table_row("fail_reason", dut, table)
        self._add_dut_table_row("comment", dut, table)

    def _add_dut_table_row(self, test_field, dut, table):
        """Create a row in the DUT result table

        Args:
            test_field (str): test parameter to add to table
            dut (dict): dictionary of test parameters
            table (obj): word doc table obj
        """

        if test_field in dut:
            test_value = dut[test_field]
            test_field = self._format_test_field(test_field)

            row_cells = table.add_row().cells
            row_cells[0].text = test_field
            row_cells[1].text = str(test_value)
        elif test_field == "test_id":
            test_value = self._test_id
            test_field = self._format_test_field(test_field)

            row_cells = table.add_row().cells
            row_cells[0].text = test_field
            row_cells[1].text = str(test_value)

            self._test_id += 1

    def _compile_suite_results(self):
        """Compile test suite results and return them

        Return:
            suite_results (list): List of compiled test suite data
        """

        logging.info(
            "The following test suites have been collected "
            f"{self._results_datamodel}"
        )
        if not self._results_datamodel:
            logging.info("Skipping the compiled test suite result")
            return
        test_suites = self._results_datamodel["test_suites"]
        suite_results = []

        for test_suite in test_suites:
            suite_result = {}
            suite_result["total_tests"] = 0
            suite_result["total_pass"] = 0
            suite_result["total_fail"] = 0
            suite_result["total_skip"] = 0
            suite_result["name"] = test_suite["name"]

            suite_name = suite_result["name"]
            logging.info(
                "Zeroing test_suite results for test suite: "
                f"{suite_name} and data: {suite_result}"
            )

            for test_case in test_suite["test_cases"]:
                for dut in test_case["duts"]:
                    suite_result["total_tests"] += 1

                    if dut["test_result"] and dut["test_result"] == "Skipped":
                        suite_result["total_skip"] += 1
                    elif dut["test_result"] and dut["test_result"] != "Skipped":
                        suite_result["total_pass"] += 1
                    else:
                        suite_result["total_fail"] += 1

            logging.info(f"Compiled test suite data: {suite_result}")
            suite_results.append(suite_result)

        logging.info(f"Compiled suite results: {suite_results}")
        return suite_results

    def _compile_testcase_results(self):
        """Compile test case results and return them"""

        if not self._results_datamodel:
            logging.info("Skipping test case results")
            return
        test_suites = self._results_datamodel["test_suites"]
        testcase_results = []

        for test_suite in test_suites:
            test_cases = test_suite["test_cases"]
            ts_name = self._format_ts_name(test_suite["name"])
            logging.info(f"Compiling results for test suite {ts_name}")

            for test_case in test_cases:
                tc_name = self._format_tc_name(test_case["name"])
                logging.info(f"Compiling results for test case {tc_name}")
                duts = test_case["duts"]

                for dut in duts:
                    testcase_result = {}

                    dut_names = dut["duts"]
                    fail_reason = dut["fail_reason"]
                    logging.info(f"Compiling results for DUT/s {dut_names}")

                    if dut["test_result"] and dut["test_result"] == "Skipped":
                        test_result = "SKIP"
                    elif dut["test_result"] and dut["test_result"] != "Skipped":
                        test_result = "PASS"
                    else:
                        test_result = "FAIL"

                    testcase_result["test_suite"] = ts_name
                    testcase_result["test_case"] = tc_name
                    testcase_result["dut"] = ", ".join(dut_names)
                    testcase_result["results"] = test_result
                    testcase_result["fail_reason"] = fail_reason
                    logging.info(f"Compiled results: {testcase_result}")

                    testcase_results.append(testcase_result)
                    logging.info(
                        "After testcase results struct appended: "
                        f"{testcase_results}"
                    )

                logging.info(
                    "Interim dut -- testcase results struct "
                    f"{testcase_results}"
                )

        logging.info(f"Returning testcase result {testcase_results}")
        return testcase_results

    def _format_ts_name(self, ts_name):
        """Input a test suite program name and return a formatted name for
            test suite

        Args:
            ts_name (str): Name of test suite program

        Return:
            ts_name (str): Formatted test suite name
        """

        logging.info(f"Test suite name is {ts_name}")
        ts_name = ts_name.split(".")[0]
        ts_name = ts_name.split("_")[1].capitalize()
        logging.info(f"Formatted test suite name is {ts_name}")

        return ts_name

    def _format_tc_name(self, tc_name):
        """Input a PyTest test case  name and return a formatted name for
            test case

        Args:
            tc_name (str): Name of PyTest test case

        Return:
            tc_name (str): Formatted test case name
        """

        logging.info(f"Test case name is {tc_name}")
        tc_name = " ".join(tc_name.split("_"))
        tc_name = tc_name.replace("intf", "interface")
        tc_name = tc_name.capitalize()
        logging.info(f"Formattted test case name is {tc_name}")

        return tc_name

    def _format_test_field(self, test_field):
        """Input a test field name and return a formatted name for
            test field

        Args:
            test_field (str): Name of test field

        Return:
            test_field (str): Formatted test field
        """

        logging.info(f"Test case name is {test_field}")
        test_field = str(test_field)
        test_field = " ".join(test_field.split("_"))
        test_field = test_field.replace("cmd", "command")
        test_field = test_field.replace("dut", "device under Test")
        test_field = test_field.capitalize()
        logging.info(f"Formattted test field is {test_field}")

        return test_field
