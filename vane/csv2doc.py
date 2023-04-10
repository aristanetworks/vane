#!/usr/bin/env python3
#
# Copyright (c) 2019, Arista Networks EOS+
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice, this
#   list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the Arista nor the names of its
#   contributors may be used to endorse or promote products derived from
#   this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE AR
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#

""" Utility script for taking a manual testing spreadsheet and converting it
    it into a formattted word doc.

    Runtime options:
        --avd: Will parse AVD NRFU CSV and add it to the test report

    Input Variables:
        TEST_SUITE_HEADER (str): Regex to match Test Case ID
        CSVFILES (list): List of path + file names for CSV input data
        AVDFILES (list): List of path + file names for AVD NRFU CSV input data
        DOCFILE (str): Path + file name for output Word doc
        EOS (str): EOS version number
        TITLE_ROW (int): Title row in CSV
        HEADER_ROW (int): Header row in CSV
        TC_ID (str): Regex to match test case ID
        CASE_NAME (str): Regex to match test case name
        PROCEDURE (str): Regex to match test steps / procedures
        EXPECTED_RESULT (str): Regex to match test criteria
        PASS_FAIL (str): Regex to match test status: pass or fail
        OBSERVATION (str): Regex to match test completion note
        TC_TYPE (str): Regex to match test case type
"""

import csv
import re
import argparse
import docx
import sys
from pprint import pprint
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_COLOR_INDEX
from docx.table import Table

# Global Variables
# test suite header
TEST_SUITE_HEADER = r"\d{1,2}\.\d{1,2}"
# Input Variables
# Name of CSV to use for test output
CSVFILES = [
    "../Arista Agora Test Plan V9.3 - EOS Test Items.csv",
    "../Arista Agora Test Plan V9.3 - CVA_CVP Test Items.csv",
]
# Name of AVD CSV files for test output
AVDFILES = ["../DC1_FABRIC-state.csv"]
DOCFILE = "../test_summary_results.docx"
# EOS version
EOS = "4.29.1F-DPE"
# Column name of interesting values
TITLE_ROW = 0
HEADER_ROW = 1
TC_ID = r".*Test Case ID.*"
CASE_NAME = r".*Test Case Description.*"
PROCEDURE = r".*Test Steps.*"
EXPECTED_RESULT = r".*Criteria.*"
PASS_FAIL = r"Pass/Fail"
OBSERVATION = r".*Observations.*"
TC_TYPE = r".*Test Case Type.*"
# Column Width
WIDTHS = (
    Inches(0.5),
    Inches(1.17),
    Inches(2.50),
    Inches(1.25),
    Inches(0.75),
    Inches(1.50),
)
TABLES = {
    "avd_test_summary": {
        "rows": 2,
        "header_row": {
            "headers": [
                "Total Tests",
                "Total Tests Passed",
                "Total Tests Failed",
            ]
        },
        "data_rows": {},
    },
    "avd_dut_summary": {
        "header_row": {
            "headers": [
                "Test Category",
                "Total Tests",
                "Tests Passed",
                "Tests Failed",
            ]
        },
        "data_rows": {},
    },
}


class CachedTable(Table):
    def __init__(self, tbl, parent):
        super(Table, self).__init__(parent)
        self._element = self._tbl = tbl
        self._cached_cells = None

    @property
    def _cells(self):
        if self._cached_cells is None:
            self._cached_cells = super(CachedTable, self)._cells
        return self._cached_cells

    @staticmethod
    def transform(table):
        cached_table = CachedTable(table._tbl, table._parent)
        return cached_table


def parse_cli():
    """Parse cli options.

    Returns:
        args (obj): An object containing the CLI arguments.
    """

    parser = argparse.ArgumentParser(description="Convert CSV files to Word Docx")

    parser.add_argument(
        "--avd",
        dest="avd",
        action="store_true",
        help="Convert AVD NRFU to Word Docx",
    )

    args = parser.parse_args()

    return args


def parse_csv(csvfiles):
    """Iterate through a list of CSV files.  Open and parse CSV data
       and find test data

    Args:
        csvfiles (list): List of CSV file names

    Returns:
        list: List of interesting test data
    """
    csv_list = []

    # iterate through CSV files
    for csvfile in csvfiles:
        print(f"  - Parsing CSV file: {csvfile} for test data...")

        # open csv
        with open(csvfile, "r", encoding="utf-8") as file:
            csv_reader = csv.reader(file)
            csv_list.extend(csv_reader)

    return csv_list


def set_header_columns(row):
    """Use regex to find a column index for interesting data

    Args:
        row (list): CSV row

    Returns:
        dict: key is column name and value is column index
    """

    header_columns = {}

    # iterate through cells to find column headers
    for count, _ in enumerate(row):
        if re.match(TC_ID, row[count]):
            header_columns["tc_id"] = count
        elif re.match(CASE_NAME, row[count]):
            header_columns["case_name"] = count
        elif re.match(PROCEDURE, row[count]):
            header_columns["procedure"] = count
        elif re.match(EXPECTED_RESULT, row[count]):
            header_columns["expected_result"] = count
        elif re.match(PASS_FAIL, row[count]):
            header_columns["pass_fail"] = count
        elif re.match(OBSERVATION, row[count]):
            header_columns["observation"] = count
        elif re.match(TC_TYPE, row[count]):
            header_columns["tc_type"] = count

    return header_columns


def inspect_test_data(csv_reader):
    """Inspect each row of CSV data and determine if its interesting

    Args:
        csv_reader (obj): CSV object to iterate through

    Returns:
        list: List of interesting test data
    """

    test_data = []

    for count, row in enumerate(csv_reader):
        test_data_entry = []

        # skip title row
        if count == TITLE_ROW:
            pass
        # inspect header and find correct column numbers
        elif count == HEADER_ROW:
            header_columns = set_header_columns(row)
        # parse csv row for a test case ids
        elif "TN" in row[header_columns["tc_id"]]:
            test_data_entry.append(row[header_columns["tc_id"]])
            test_data_entry.append(row[header_columns["case_name"]])
            test_data_entry.append(row[header_columns["procedure"]])
            test_data_entry.append(row[header_columns["expected_result"]])
            test_data_entry.append(row[header_columns["pass_fail"]])
            test_data_entry.append(row[header_columns["observation"]])
            test_data.append(test_data_entry)
        # parse csv row for header
        elif re.match(TEST_SUITE_HEADER, row[header_columns["case_name"]]):
            test_data_entry.append(row[header_columns["case_name"]])
            test_data.append(test_data_entry)
        # parse csv row for sub-section
        elif (not row[header_columns["tc_id"]] and not row[header_columns["tc_type"]]) and row[
            header_columns["case_name"]
        ] != "":
            test_data_entry.append("")
            test_data_entry.append(row[header_columns["case_name"]])
            test_data.append(test_data_entry)

    return test_data


def inspect_avd_data(csv_reader):
    """Parses AVD NRFU CSV.  Creates mutliple data structs for outputing results.

    Args:
        csv_reader (obj): CSV object to iterate through

    Returns:
        dict: Calculated total test summary data
        dict: Calculated summary data by DUTs
        dict: Calculated test summary data by test case categories
    """
    # initialize counters
    tests_summary = {
        "test_total": 0,
        "pass_counter": 0,
        "fail_counter": 0,
    }
    dut_summary = {}
    cat_summary = {}

    for count, row in enumerate(csv_reader):
        # skip title row
        if count == TITLE_ROW:
            pass
        else:
            # parse and build data structs for tables
            tests_summary = calc_summary_results(row, tests_summary)
            dut_summary = calc_dut_results(row, dut_summary)
            cat_summary = calc_cat_results(row, cat_summary)

    pprint(dut_summary)
    sys.exit()

    return tests_summary, dut_summary, cat_summary


def calc_summary_results(row, tests_summary):
    """Does total test summary caluclations for 1 row of NRFU Data and returns an updated total

    Args:
        row (list): CSV NRFU test cases entry
        tests_summary (dict): Ongoing calculated total test summary data

    Returns:
        dict: Ongoing calculated total test summary data
    """
    # calculate total pass/fail
    if row[5] == "PASS":
        tests_summary["pass_counter"] += 1
    elif row[5] == "FAIL":
        tests_summary["fail_counter"] += 1
    else:
        print("Bogus information detected: {row}")

    tests_summary["test_total"] += 1

    return tests_summary


def calc_dut_results(row, dut_summary):
    """Does total per DUT test summary calcuations for 1 row of NRFU Data returns an updated total

    Args:
        row (list): CSV NRFU test cases entry
        dut_summary (dict): Ongoing calculated summary data by DUTs

    Returns:
        dict: Ongoing calculated summary data by DUTs
    """
    dut = row[1]

    if dut not in dut_summary:
        # initialize dut counters
        dut_summary[dut] = {
            "test_total": 0,
            "pass_counter": 0,
            "fail_counter": 0,
            "test_category": [],
        }
    # calculate dut total pass/fail
    if row[5] == "PASS":
        dut_summary[dut]["pass_counter"] += 1
    elif row[5] == "FAIL":
        dut_summary[dut]["fail_counter"] += 1
        if row[2] not in dut_summary[dut]["test_category"]:
            dut_summary[dut]["test_category"].append(row[2])

    dut_summary[dut]["test_total"] += 1

    return dut_summary


def calc_cat_results(row, cat_summary):
    """Does total per category test summary calculations for 1 row of NRFU data returns an
       updated total

    Args:
        row (list): CSV NRFU test cases entry
        cat_summary (dict): Ongoing calculated summary data by category

    Returns:
        dict: Ongoing calculated summary data by category
    """
    category = row[2]

    if category not in cat_summary:
        # initialize category counters
        cat_summary[category] = {
            "test_total": 0,
            "pass_counter": 0,
            "fail_counter": 0,
        }

    # calculate category total pass/fail
    if row[5] == "PASS":
        cat_summary[category]["pass_counter"] += 1
    elif row[5] == "FAIL":
        cat_summary[category]["fail_counter"] += 1

    cat_summary[category]["test_total"] += 1

    return cat_summary


def write_report():
    """Creates report and sets initial doc parameters

    Returns:
        obj: Object representing Word Docx
    """
    print(f"Writing report doc: {DOCFILE}")

    # Open doc
    document = docx.Document()
    # Change doc margins
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    return document


def write_avd_test_summary(document, test_data, tests_summary, dut_summary, cat_summary):
    """Creates AVD Test summary sections within report

    Args:
        document (obj): Test report Word doc
        test_data (list): AVD NRFU test case results
        tests_summary (dict): Calculated total test summary data
        dut_summary (dict): Calculated summary data by DUTs
        cat_summary (dict): Calculated test summary data by test case categories
    """
    print("  - Writing AVD test case summary section")

    header_text = "5	AVD NRFU Validate State Report"
    write_header(document, header_text)

    doc_text = "\nTest Results Summary"
    write_text(document, doc_text, font_pt=14, bold=True)

    write_table(document, TABLES["avd_test_summary"], tests_summary, "Summary Totals")
    write_dut_summary(document, dut_summary)
    write_category_summary(document, cat_summary)

    failed_tests = [v for v in test_data if "FAIL" in v[5]]
    write_testcase_summary(document, failed_tests, "Failed Test Cass Results Summary")
    write_testcase_summary(document, test_data, "Test Cases Results Summary")


def write_testcase_summary(document, tests_data, doc_text):
    """Writes detailed summary of each AVD NRFU test case

    Args:
        document (obj): Test report Word doc
        test_data (list): AVD NRFU test case results
        doc_text (str): Header name for the section
    """
    print(f"    * Writing {doc_text} Section...")
    write_text(document, f"\n{doc_text}", font_pt=14, bold=True)

    headers = [
        "Test ID",
        "Node",
        "Test Category",
        "Test Description",
        "Test",
        "Test Result",
        "Failure Reason",
    ]
    columns = len(headers)
    rows = len(tests_data)

    table = CachedTable.transform(document.add_table(rows=rows, cols=columns, style="Table Grid"))
    write_header_row(table, headers)

    row = 1
    for test_case_data in tests_data:
        if row != HEADER_ROW:
            for column, test_data in enumerate(test_case_data):
                write_cell(table, str(test_data), column, (row - 1))
        row += 1


def write_table(document, table_parameters, tests_parameters, doc_text):
    print(f"    * Writing table: {doc_text}..")
    doc_text = f"\n{doc_text}"
    write_text(document, doc_text, font_pt=12, bold=True)

    headers = table_parameters["header_row"]["headers"]
    columns = len(headers)
    rows = table_parameters["rows"]

    table = CachedTable.transform(document.add_table(rows=rows, cols=columns, style="Table Grid"))

    for column, header in enumerate(headers):
        write_cell(table, header, column, 0, font="Arial", font_pt=9, bold=True, color="00FFFF")

    for column, tests_parameter in enumerate(tests_parameters):
        write_cell(table, str(tests_parameters[tests_parameter]), column, 1)


def write_dut_summary(document, dut_summary):
    """Writes test case summary per dut for AVD test cases

    Args:
        document (obj): Test report Word doc
        dut_summary (dict): Calculated summary data by DUTs
    """
    print("    * Writing DUT Summary Section...")
    doc_text = "\nSummary Totals Device Under Tests"
    write_text(document, doc_text, font_pt=12, bold=True)

    # create table
    table = document.add_table(rows=1, cols=5, style="Table Grid")
    table.autofit = False

    # create table rows
    headers = ["DUT", "Total Tests", "Tests Passed", "Tests Failed", "Catorgories Failed"]
    write_header_row(table, headers)

    for row, dut in enumerate(dut_summary):
        _ = table.add_row().cells
        write_cell(table, dut, 0, (row + 1))
        for column, test_data in enumerate(dut_summary[dut]):
            if test_data == "test_category":
                dut_summary[dut][test_data] = ", ".join(dut_summary[dut][test_data])

            write_cell(table, str(dut_summary[dut][test_data]), (column + 1), (row + 1))


def write_category_summary(document, cat_summary):
    """Writes test case  summary per category for AVD test cases

    Args:
        document (obj): Test report Word doc
        cat_summary (dict): Calculated test summary data by test case categories
    """
    print("    * Writing Category Summary Section...")
    doc_text = "\nSummary Totals Per Category"
    write_text(document, doc_text, font_pt=12, bold=True)

    # create table
    table = document.add_table(rows=1, cols=4, style="Table Grid")
    table.autofit = False

    # create table rows
    headers = [
        "Test Category",
        "Total Tests",
        "Tests Passed",
        "Tests Failed",
    ]
    write_header_row(table, headers)

    for row, dut in enumerate(cat_summary):
        _ = table.add_row().cells
        write_cell(table, dut, 0, (row + 1))
        for column, test_data in enumerate(cat_summary[dut]):
            write_cell(table, str(cat_summary[dut][test_data]), (column + 1), (row + 1))


# pylint: disable-next=too-many-arguments
def write_cell(
    table,
    text,
    column,
    row,
    font=None,
    font_pt=10,
    bold=False,
    color=None,
    text_color=None,
):
    """Writes a cell within Word doc table

    Args:
        table (obj): Word doc obj representing a table
        text (str): Text to output in table cell
        column (int): Column number in table cell
        row (int): Row number in table cell
        font (str, optional): Font to use in table cell. Defaults to None.
        font_pt (int, optional): Font size to use in table cell. Defaults to None.
        bold (bool, optional): Bold text in table cell. Defaults to False.
        color (str, optional): Hex-decimanal color to fill table cell. Defaults to None.
        format (str, optional): Style of outputting text in table cell. Defaults to "string".
    """
    para = table.cell(row, column).paragraphs[0]
    write_text(None, text, font=font, font_pt=font_pt, bold=bold, color=text_color, para=para)

    if color:
        shade_cell(column, table, row, color)


def write_manual_test_summary(document, test_data):
    """Writes detailed summary of each manual test case

    Args:
        document (obj): Test report Word doc
        test_data (list): List of interesting test data
    """
    print("  - Writing manual test case summary section")

    header_text = "4	Test Summary Result"
    write_header(document, header_text)

    doc_text = (
        f"\nTest Cases for Arista BAU Hardware & EOS {EOS} Certification. "
        "Please note – for more details on test outputs please reference "
        "the test case output appendix docs.\n\n\n"
    )
    write_text(document, doc_text)
    write_results_table(document, test_data)
    doc_text = (
        "\n\n\nPlease note – for more details on test outputs please "
        "reference the test case output appendix docs.\n"
    )
    write_text(document, doc_text)


def write_header(document, header_text):
    """Write test section header to word doc

    Args:
        document (obj): Test report Word doc
        header_text (list): Header description field
    """

    # Write summary header
    heading = document.add_heading(level=2)
    run = heading.add_run(header_text)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(204, 0, 0)


# pylint: disable-next=too-many-arguments
def write_text(document, doc_text, font_pt=10, font="Arial", bold=False, para=None, color=None):
    """Write a paragraph to Word Doc

    Args:
        document (obj): Word doc object
        doc_text (str): Text to write to doc
        font_pt (int): Word font size
        font (string): Word font style
        bold (bool): Bold text
        para (obj): Word doc text object
        color (obj): Word font color
    """
    if not para:
        para = document.add_paragraph()

    run = para.add_run(doc_text)
    run.font.size = Pt(font_pt)
    run.font.name = font

    if bold:
        run.font.bold = bold

    if color:
        run.font.color.rgb = color


def write_header_row(table, headers, width=False):
    """Write the header row for the result table

    Args:
        table (obj): Word doc table object
        headers (list): Table header fields
        width (bool): Width of a cell
    """

    # Write header row in table
    for column, header in enumerate(headers):
        if width:
            set_cell_widths(table, 0)

        write_cell(table, header, column, 0, font="Arial", font_pt=9, bold=True, color="00FFFF")


def write_results_table(document, test_data):
    """Write results table to Word doc

    Args:
        document (obj): Word doc object
        test_data (list): List of interesting test data
    """

    # create table
    table = document.add_table(rows=1, cols=6, style="Table Grid")
    table.autofit = False

    # create table rows
    headers = [
        "Test Case",
        "Case Name",
        "Procedure",
        "Expected result",
        "Pass/Fail",
        "Observation",
    ]
    write_header_row(table, headers, width=True)
    write_data_row(table, test_data)


def write_data_row(table, test_data):
    """Write a test case result row in table

    Args:
        table (obj): Word doc table object
        test_data (list): List of interesting test data
    """

    # Iterate through test data and write table rows
    for count, test_case_entry in enumerate(test_data):
        count += 1
        row_cells = table.add_row().cells
        set_cell_widths(table, count)

        # Identify test suite sub-header row
        if re.match(TEST_SUITE_HEADER, test_case_entry[0]):
            row_cells[0].merge(row_cells[5])
            row_cells[0].text = ""
            run = row_cells[0].paragraphs[0].add_run(test_case_entry)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.bold = True
            shade_cell(0, table, count, "CCFFFF")

        # Identify test case result data row
        elif "TN" in test_case_entry[0]:
            for counter, test_data_entry in enumerate(test_case_entry):
                row_cells[counter].text = ""
                run.font.name = "Arial"
                run.font.size = Pt(9)

                # Test Case Identifier cell
                if counter == 0:
                    test_data_entry = format_test_id(test_data_entry)
                    run = row_cells[counter].paragraphs[0].add_run(test_data_entry)
                # Test Case Pass/Fail cell
                elif counter == 4:
                    format_pass_fail(row_cells, counter, test_data_entry)
                # Cells not needing special formatting
                else:
                    run = row_cells[counter].paragraphs[0].add_run(test_data_entry)

                run.font.name = "Arial"
                run.font.size = Pt(9)

        # identify test suite sub section
        elif test_case_entry[0] == "":
            row_cells[2].merge(row_cells[5])

            run = row_cells[1].paragraphs[0].add_run(test_case_entry[1])
            run.font.name = "Arial"
            run.font.size = Pt(9)

            for cell_idx in [0, 1, 2]:
                shade_cell(cell_idx, table, count, "CCFFFF")


def shade_cell(cell_idx, table, count, shade):
    """Shade a cell in word doc table

    Args:
        cell_idx (int): Column index for cell to shade
        table (obj): Word doc table object
        count (int): Row index for cell to shade
        shade (str): hexadecimal color representation
    """
    cell = table.cell(count, cell_idx)
    color = parse_xml(
        # pylint: disable-next=consider-using-f-string
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), shade)
    )
    # pylint: disable-next=protected-access
    cell._tc.get_or_add_tcPr().append(color)


def format_test_id(test_data_entry):
    """Format the test case id

    Args:
        test_data_entry (str): Text for table cell

    Returns:
        str: Formatted text for table cell
    """
    if "TN" in test_data_entry:
        test_data_entry = test_data_entry.split("TN")[1]

    return test_data_entry


def format_pass_fail(row_cells, counter, test_data_entry):
    """Format pass fail cell

    Args:
        row_cells (obj): Table cell
        counter (int): Table cell number
        test_data_entry (str): Text for table cell
    """
    run = row_cells[counter].paragraphs[0].add_run(test_data_entry.upper())

    if test_data_entry.upper() == "PASS":
        # pylint: disable-next=no-member
        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        run.font.bold = True
    elif test_data_entry.upper() == "FAIL":
        # pylint: disable-next=no-member
        run.font.highlight_color = WD_COLOR_INDEX.RED
        run.font.bold = True
    else:
        # pylint: disable-next=no-member
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.bold = True


def set_cell_widths(table, row):
    """Set width of Word table cell

    Args:
        table (obj): Word doc table object
        row (int): Table row number for operations
    """

    for column, width in enumerate(WIDTHS):
        table.rows[row].cells[column].width = width


def main():
    """Main Python function"""

    print("Starting...")
    # Input CLI arguments
    args = parse_cli()

    # parse csv
    print("Parsing CSV files:")
    csv_list = parse_csv(CSVFILES)
    test_data = inspect_test_data(csv_list)

    if args.avd:
        avd_data = parse_csv(AVDFILES)
        tests_summary, dut_summary, cat_summary = inspect_avd_data(avd_data)

    # write word doc
    document = write_report()
    write_manual_test_summary(document, test_data)

    if args.avd:
        write_avd_test_summary(document, avd_data, tests_summary, dut_summary, cat_summary)

    # Close doc
    document.save(DOCFILE)
    print("Complete!")


if __name__ == "__main__":
    main()
